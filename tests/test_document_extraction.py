"""Tests for document extraction in bot_media."""

from __future__ import annotations

import json
import unittest
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pymupdf
import pytest

from analyst.delivery.bot_media import (
    RequestDocumentInput,
    _document_file_ref,
    _extract_attached_document,
    _extract_text_from_bytes,
    _summarize_user_message,
)
from analyst.delivery.bot_constants import (
    MAX_DOCUMENT_FILE_SIZE,
    MAX_DOCUMENT_TEXT_CHARS,
)


# ---------------------------------------------------------------------------
# _document_file_ref
# ---------------------------------------------------------------------------


class TestDocumentFileRef:
    def test_supported_pdf(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="abc123",
                mime_type="application/pdf",
                file_size=5000,
                file_name="report.pdf",
            )
        )
        result = _document_file_ref(msg)
        assert result == ("abc123", "application/pdf", "report.pdf", 5000)

    def test_supported_text(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="txt1",
                mime_type="text/plain",
                file_size=100,
                file_name="notes.txt",
            )
        )
        result = _document_file_ref(msg)
        assert result is not None
        assert result[1] == "text/plain"

    def test_unsupported_mime(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="doc1",
                mime_type="application/zip",
                file_size=100,
                file_name="file.zip",
            )
        )
        assert _document_file_ref(msg) is None

    def test_supported_docx(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="docx1",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=5000,
                file_name="report.docx",
            )
        )
        result = _document_file_ref(msg)
        assert result is not None
        assert result[2] == "report.docx"

    def test_supported_xlsx(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="xlsx1",
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                file_size=3000,
                file_name="data.xlsx",
            )
        )
        result = _document_file_ref(msg)
        assert result is not None
        assert result[2] == "data.xlsx"

    def test_too_large(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="big1",
                mime_type="application/pdf",
                file_size=MAX_DOCUMENT_FILE_SIZE + 1,
                file_name="huge.pdf",
            )
        )
        assert _document_file_ref(msg) is None

    def test_no_document(self):
        msg = SimpleNamespace(document=None)
        assert _document_file_ref(msg) is None

    def test_missing_filename_gets_default(self):
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="nf1",
                mime_type="text/plain",
                file_size=10,
                file_name="",
            )
        )
        result = _document_file_ref(msg)
        assert result is not None
        assert result[2] == "document"

    def test_image_mime_rejected(self):
        """Image documents should NOT be handled by document extractor."""
        msg = SimpleNamespace(
            document=SimpleNamespace(
                file_id="img1",
                mime_type="image/jpeg",
                file_size=1000,
                file_name="photo.jpg",
            )
        )
        assert _document_file_ref(msg) is None


# ---------------------------------------------------------------------------
# _extract_text_from_bytes
# ---------------------------------------------------------------------------


class TestExtractTextFromBytes:
    def test_plain_text_utf8(self):
        raw = "Hello, world! 你好".encode("utf-8")
        text, truncated = _extract_text_from_bytes(raw, "text/plain", "file.txt")
        assert text == "Hello, world! 你好"
        assert truncated is False

    def test_plain_text_latin1_fallback(self):
        raw = "café résumé".encode("latin-1")
        text, truncated = _extract_text_from_bytes(raw, "text/plain", "file.txt")
        assert "caf" in text
        assert truncated is False

    def test_json_file(self):
        data = json.dumps({"key": "value"}).encode("utf-8")
        text, truncated = _extract_text_from_bytes(data, "application/json", "data.json")
        assert '"key"' in text
        assert truncated is False

    def test_truncation(self):
        raw = ("A" * (MAX_DOCUMENT_TEXT_CHARS + 500)).encode("utf-8")
        text, truncated = _extract_text_from_bytes(raw, "text/plain", "big.txt")
        assert len(text) == MAX_DOCUMENT_TEXT_CHARS
        assert truncated is True

    def test_pdf_extraction(self):
        # Create a small PDF with pymupdf
        doc = pymupdf.open()
        page = doc.new_page(width=200, height=200)
        page.insert_text((50, 100), "Test PDF content")
        pdf_bytes = doc.tobytes()
        doc.close()

        text, truncated = _extract_text_from_bytes(pdf_bytes, "application/pdf", "test.pdf")
        assert "Test PDF content" in text
        assert truncated is False

    def test_csv_file(self):
        raw = "name,age\nAlice,30\nBob,25\n".encode("utf-8")
        text, truncated = _extract_text_from_bytes(raw, "text/csv", "data.csv")
        assert "Alice" in text
        assert truncated is False

    def test_docx_extraction(self):
        from docx import Document as DocxDocument
        from io import BytesIO

        doc = DocxDocument()
        doc.add_paragraph("Hello from Word")
        doc.add_paragraph("Second paragraph")
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        text, truncated = _extract_text_from_bytes(docx_bytes, mime, "test.docx")
        assert "Hello from Word" in text
        assert "Second paragraph" in text
        assert truncated is False

    def test_xlsx_extraction(self):
        from openpyxl import Workbook
        from io import BytesIO

        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Age"])
        ws.append(["Alice", 30])
        ws.append(["Bob", 25])
        buf = BytesIO()
        wb.save(buf)
        xlsx_bytes = buf.getvalue()

        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        text, truncated = _extract_text_from_bytes(xlsx_bytes, mime, "test.xlsx")
        assert "[Sheet: Data]" in text
        assert "Alice" in text
        assert "30" in text
        assert truncated is False


# ---------------------------------------------------------------------------
# _summarize_user_message with document
# ---------------------------------------------------------------------------


class TestSummarizeWithDocument:
    def test_text_and_document(self):
        doc = RequestDocumentInput(
            text="content", filename="report.pdf", mime_type="application/pdf"
        )
        result = _summarize_user_message("Check this", document=doc)
        assert "[Document: report.pdf]" in result
        assert result.startswith("Check this")

    def test_document_only(self):
        doc = RequestDocumentInput(
            text="content", filename="notes.txt", mime_type="text/plain"
        )
        result = _summarize_user_message("", document=doc)
        assert result == "[Document: notes.txt]"

    def test_no_document(self):
        result = _summarize_user_message("Hello")
        assert result == "Hello"


# ---------------------------------------------------------------------------
# _extract_attached_document (async, mocked Telegram API)
# ---------------------------------------------------------------------------


def _make_telegram_document_message(
    *,
    file_id: str = "file123",
    mime_type: str = "text/plain",
    file_size: int = 100,
    file_name: str = "test.txt",
):
    return SimpleNamespace(
        document=SimpleNamespace(
            file_id=file_id,
            mime_type=mime_type,
            file_size=file_size,
            file_name=file_name,
        ),
        photo=None,
        reply_to_message=None,
    )


class TestExtractAttachedDocumentAsync(unittest.IsolatedAsyncioTestCase):
    async def test_direct_message(self):
        raw = b"Hello from document"
        telegram_file = AsyncMock()
        telegram_file.download_as_bytearray = AsyncMock(return_value=bytearray(raw))

        bot = AsyncMock()
        bot.get_file = AsyncMock(return_value=telegram_file)
        ctx = SimpleNamespace(bot=bot)

        message = _make_telegram_document_message()
        update = SimpleNamespace(effective_message=message)

        result = await _extract_attached_document(update, ctx)
        assert result is not None
        assert result.text == "Hello from document"
        assert result.filename == "test.txt"
        assert result.source == "message"
        assert result.truncated is False

    async def test_from_reply(self):
        raw = b"Reply doc text"
        telegram_file = AsyncMock()
        telegram_file.download_as_bytearray = AsyncMock(return_value=bytearray(raw))

        bot = AsyncMock()
        bot.get_file = AsyncMock(return_value=telegram_file)
        ctx = SimpleNamespace(bot=bot)

        reply_msg = _make_telegram_document_message(file_name="reply.txt")
        message = SimpleNamespace(
            document=None,
            photo=None,
            reply_to_message=reply_msg,
        )
        update = SimpleNamespace(effective_message=message)

        result = await _extract_attached_document(update, ctx)
        assert result is not None
        assert result.source == "reply"
        assert result.filename == "reply.txt"

    async def test_none_when_unsupported(self):
        message = SimpleNamespace(
            document=SimpleNamespace(
                file_id="x",
                mime_type="application/zip",
                file_size=100,
                file_name="archive.zip",
            ),
            photo=None,
            reply_to_message=None,
        )
        update = SimpleNamespace(effective_message=message)
        ctx = SimpleNamespace(bot=AsyncMock())

        result = await _extract_attached_document(update, ctx)
        assert result is None
